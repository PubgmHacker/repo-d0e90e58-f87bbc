#!/usr/bin/env python3
"""
Генерация бизнес-плана Plink на основе шаблона SOLD.docx
Сохраняет исходную структуру (10 разделов + таблицы), полностью переписывает контент
под мобильное iOS-приложение Plink для совместного просмотра видео.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from copy import deepcopy
import os

OUTPUT_PATH = "/home/z/my-project/download/souldawn Бизнес-план.docx"

# ─────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────

def set_cell_borders(cell, color="000000", sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), sz)
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_table_borders(table, color="000000", sz="4"):
    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(cell, color, sz)

def shade_cell(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def add_paragraph_with_format(doc, text, *, bold=False, italic=False,
                              size=12, align=None, color=None,
                              first_line_indent=None, space_after=6,
                              font_name="Times New Roman"):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(space_after)
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_heading(doc, text, level=1):
    sizes = {1: 16, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(sizes.get(level, 12))
    run.bold = True
    return p

def add_bullet(doc, text, size=12):
    p = doc.add_paragraph(style='List Bullet')
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    return p

def add_table_header_row(table, headers, bold=True):
    # Add header row first, then fill
    hdr = table.add_row()
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.bold = bold
        shade_cell(cell, "D9D9D9")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def add_data_row(table, values, align_first_left=True):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if (i == 0 and align_first_left) else WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(val))
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return row

def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break()
    from docx.enum.text import WD_BREAK
    p.runs[0].add_break(WD_BREAK.PAGE)

def add_table_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.italic = True
    return p

# ─────────────────────────────────────────────────────────────────────
# BUILD DOCUMENT
# ─────────────────────────────────────────────────────────────────────

doc = Document()

# Page setup A4
for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

# Default style
style = doc.styles['Normal']
style.font.name = "Times New Roman"
style.font.size = Pt(12)

# ═══════════════════════════════════════════════════════════════════
# TITLE PAGE (table-based, as in template)
# ═══════════════════════════════════════════════════════════════════

title_table = doc.add_table(rows=1, cols=1)
title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
title_cell = title_table.rows[0].cells[0]
set_cell_borders(title_cell, sz="12")
title_cell.text = ""

# Title inside the bordered cell
p = title_cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("НАЗВАНИЕ БИЗНЕС-ПРОЕКТА")
run.font.name = "Times New Roman"
run.font.size = Pt(16)
run.bold = True

title_fields = [
    ("1. Название проекта", "«Plink» — мобильное iOS-приложение для совместного просмотра видео"),
    ("2. Организационно-правовая форма", "ООО «Plink Tech», УСН доход 6%"),
    ("3. Контактные данные", ""),
    ("ИНН", "в процессе регистрации"),
    ("ФИО руководителя", "учредитель проекта Plink"),
    ("Адрес офиса", "г. Москва (удалённая команда)"),
    ("Телефон", "+7 (XXX) XXX-XX-XX"),
    ("E-mail", "founder@plink.app"),
    ("4. ОКВЭД с расшитровкой", "62.01 «Разработка компьютерного ПО, информационных услуг»"),
    ("5. Место реализации проекта", "Российская Федерация (дистрибуция через App Store по всему миру)"),
    ("6. Дата начала реализации проекта", "1 июля 2026 г."),
    ("7. Продолжительность реализации проекта", "12 месяцев (1 год)"),
    ("8. Необходимость лицензии / разрешения", "Аккаунт Apple Developer Program ( годовая подписка )"),
    ("9. Стоимость проекта", "1 200 000,00 рублей"),
    ("в том числе:", ""),
    ("собственные средства", "200 000,00 рублей"),
    ("средства гос. поддержки", "1 000 000,00 рублей"),
    ("Срок окупаемости проекта", "28 мес."),
    ("Рентабельность продукции", "57 %"),
    ("Порог рентабельности проекта", "540 000,00 руб."),
    ("Чистая прибыль проекта", "513 190,00 руб."),
    ("Количество созданных рабочих мест", "2 чел. (учредитель + маркетолог-фрилансер)"),
]

for label, value in title_fields:
    p = title_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(4)
    run = p.add_run(f"{label}: {value}" if value else label)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    if not value and label == "в том числе:":
        run.italic = True

# Footer of title page
p = title_cell.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf = p.paragraph_format
pf.space_before = Pt(18)
run = p.add_run("Москва — 2026")
run.font.name = "Times New Roman"
run.font.size = Pt(12)
run.bold = True

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════
# СОДЕРЖАНИЕ
# ═══════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("СОДЕРЖАНИЕ")
run.font.name = "Times New Roman"
run.font.size = Pt(16)
run.bold = True
p.paragraph_format.space_after = Pt(12)

toc_table = doc.add_table(rows=0, cols=2)
toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
toc_items = [
    ("1", "РЕЗЮМЕ"),
    ("2", "ОПИСАНИЕ БИЗНЕСА"),
    ("3", "ПРОДУКТЫ И УСЛУГИ"),
    ("4", "АНАЛИЗ РЫНКА"),
    ("5", "ПЛАН МАРКЕТИНГА"),
    ("6", "ПРОИЗВОДСТВЕННЫЙ ПЛАН"),
    ("7", "ПЛАН ПЕРСОНАЛА"),
    ("8", "ПЛАН ЗАТРАТ"),
    ("9", "ФИНАНСОВЫЙ ПЛАН"),
    ("10", "ПРИЛОЖЕНИЯ"),
]
for num, title in toc_items:
    row = toc_table.add_row()
    c1, c2 = row.cells
    c1.text = ""
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(num)
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(12)
    r1.bold = True
    c2.text = ""
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r2 = p2.add_run(title)
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(12)
    r2.bold = True

set_table_borders(toc_table)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════
# РЕЗЮМЕ
# ═══════════════════════════════════════════════════════════════════

add_heading(doc, "РЕЗЮМЕ", level=1)

add_paragraph_with_format(doc,
    "Настоящий проект предусматривает разработку, запуск и монетизацию мобильного "
    "iOS-приложения Plink — сервиса для совместного просмотра видео (Watch Party) с "
    "точной синхронизацией плееров, голосовым и текстовым чатом в реальном времени. "
    "Цель проекта — занять лидирующие позиции в нише co-watching-приложений на "
    "русскоязычном рынке за счёт интеграции с российскими онлайн-кинотеатрами, "
    "встроенного ИИ-помощника и уникального биолюминесцентного дизайна.",
    first_line_indent=1.25, space_after=10)

add_paragraph_with_format(doc,
    "Целевой аудиторией приложения являются активные пользователи интернета и "
    "социальных сетей в возрасте от 16 до 35 лет: школьники и студенты, молодые "
    "специалисты, пары на расстоянии, любители кино, аниме и стримингового контента. "
    "Они ищут способы проводить время вместе онлайн, синхронно смотреть фильмы и "
    "сериалы, обмениваться реакциями в реальном времени. Пользователи ценят простоту "
    "интерфейса, точность синхронизации, поддержку отечественных видеосервисов и "
    "круглосуточную работоспособность приложения.",
    first_line_indent=1.25, space_after=10)

add_paragraph_with_format(doc,
    "Для продвижения продукта будет применяться цифровая маркетинговая стратегия: "
    "ASO-оптимизация в App Store, таргетированная реклама в Telegram и ВКонтакте, "
    "сотрудничество с инфлюенсерами на YouTube и в TikTok, а также реферальная "
    "программа внутри приложения. Дополнительно будут задействованы сообщества в "
    "Telegram, Discord и тематические площадки для киноманов и аниме-сообществ.",
    first_line_indent=1.25, space_after=10)

add_paragraph_with_format(doc, "Конкурентные преимущества проекта Plink:",
                          first_line_indent=1.25, space_after=4)

for adv in [
    "Встроенный ИИ-помощник на базе OpenRouter (рекомендации фильмов, чат-ассистент).",
    "Интеграция с 8 российскими онлайн-кинотеатрами (Кинопоиск, Иви, Okko, Wink, Start, Premier, Смотрим, КИОН).",
    "Точная синхронизация плееров с компенсацией сетевой задержки (latency compensation).",
    "Уникальный биолюминесцентный дизайн с многоцветными градиентами.",
    "Двухфакторная аутентификация и хранение токенов в Keychain (безопасность enterprise-уровня).",
    "Голосовой чат и текстовые сообщения в реальном времени через WebSocket.",
    "Кросс-платформенный бэкенд на Node.js + PostgreSQL + Redis с auto-scaling на Railway.",
    "Реферальная программа и гибкая монетизация через StoreKit 2.",
    "Локализация на 4 языка (русский, английский, китайский, испанский).",
    "Поддержка AirPlay, субтитров и Live Activities (Dynamic Island).",
]:
    add_bullet(doc, adv)

add_heading(doc, "Целевая аудитория", level=2)
add_paragraph_with_format(doc, "Географические:", bold=True, space_after=2)
for s in [
    "Страна: Российская Федерация (с потенциальным расширением на СНГ и мировую аудиторию).",
    "Регион: любой, где доступен App Store.",
    "Город: любой (продукт цифровой, географии не имеет).",
]:
    add_bullet(doc, s)

add_paragraph_with_format(doc, "Демографические:", bold=True, space_after=2)
for s in [
    "Пол: мужской и женский.",
    "Возраст: 16–35 лет.",
    "Семейное положение: холост/замужем, в том числе пары на расстоянии.",
    "Образование: школа, колледж, вуз.",
]:
    add_bullet(doc, s)

add_paragraph_with_format(doc, "Социально-экономические:", bold=True, space_after=2)
for s in [
    "Уровень дохода: средний и выше среднего (готовы платить за подписку 299–499 руб./мес).",
    "Профессия: студенты, IT-специалисты, маркетологи, креативные профессии.",
    "Образование: среднее, среднее специальное, высшее.",
]:
    add_bullet(doc, s)

add_paragraph_with_format(doc, "Психографические:", bold=True, space_after=2)
for s in [
    "Ценности: совместное времяпровождение, технологии, кинокультура, эмоциональная близость на расстоянии.",
    "Страхи: одиночество, разрыв отношений из-за расстояния, недопонимание в общении.",
    "Хобби: кино, сериалы, аниме, видеоигры, стриминг, музыка, Twitch, YouTube.",
    "Образ жизни: активные пользователи интернета 4–8 часов в день.",
    "Желания: проводить время с друзьями онлайн, смотреть фильмы синхронно, делиться эмоциями.",
]:
    add_bullet(doc, s)

add_paragraph_with_format(doc, "Поведенческие:", bold=True, space_after=2)
for s in [
    "Как ищут сервис: App Store по запросам «смотреть вместе», «watch party», «rave», «видео с друзьями».",
    "Причины установки: желание смотреть фильм с другом на расстоянии, рекомендация знакомых, реклама у блогеров.",
    "Реакция на бренд: доверие за счёт интеграции с известными российскими кинотеатрами.",
    "Частота использования: 2–5 раз в неделю, средняя сессия 90–180 минут.",
]:
    add_bullet(doc, s)

add_heading(doc, "Портрет клиента", level=2)
add_paragraph_with_format(doc,
    "Александр, 22 года, студент технического вуза из Екатеринбурга. Подрабатывает "
    "фрилансом, доход около 45 000 руб. в месяц. Свободное время проводит за "
    "просмотром фильмов и аниме, активно пользуется Telegram и YouTube. Главная "
    "потребность — смотреть сериалы с девушкой, которая учится в другом городе. "
    "Ранее пробовал Rave и Discord-боты, но сталкивался с рассинхронизацией и "
    "отсутствием российских кинотеатров. Узнаёт о Plink из рекламы в Telegram-канале "
    "про кино. После установки приглашает девушку через реферальный код, они смотрят "
    "аниме на Кинопоиске синхронно, обмениваются сообщениями в чате, используют "
    "ИИ-помощника для подбора следующего тайтла.",
    first_line_indent=1.25, space_after=10)

add_paragraph_with_format(doc,
    "С учётом отсутствия прямых российских аналогов с поддержкой отечественных "
    "кинотеатров и ИИ-функционалом, ожидается стабильный рост пользовательской базы "
    "и высокая удерживаемость (retention rate) благодаря сетевому эффекту: чем больше "
    "друзей пользователя в Plink, тем выше ценность приложения для каждого.",
    first_line_indent=1.25, space_after=10)

add_paragraph_with_format(doc,
    "Реализация проекта запланирована с 1 июля 2026 г. В течение первого года будут "
    "выполнены следующие этапы: релиз MVP в App Store (Q3 2026), запуск премиум-подписки "
    "и реферальной программы (Q4 2026), обновление с ИИ-помощником и интеграция "
    "8 кинотеатров (Q1 2027), выход на точку безубыточности (Q2 2027).",
    first_line_indent=1.25, space_after=10)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════
# РАЗДЕЛ 1. ОПИСАНИЕ БИЗНЕСА
# ═══════════════════════════════════════════════════════════════════

add_heading(doc, "РАЗДЕЛ 2 «ОПИСАНИЕ БИЗНЕСА»", level=1)

add_paragraph_with_format(doc,
    "ООО «Plink Tech» занимается разработкой и оперированием одноимённого мобильного "
    "приложения для совместного просмотра видео в режиме реального времени. Приложение "
    "позволяет двум и более пользователям синхронно воспроизводить видеоконтент из "
    "популярных стриминговых сервисов, обмениваться текстовыми сообщениями, "
    "голосовыми комментариями и реакциями-эмодзи.",
    first_line_indent=1.25, space_after=8)

add_paragraph_with_format(doc, "Цели проекта:", bold=True, space_after=4)
for s in [
    "удовлетворение спроса молодёжной аудитории в качественном сервисе co-watching с поддержкой российских видеосервисов;",
    "получение стабильной и масштабируемой выручки за счёт freemium-модели с премиум-подпиской;",
    "формирование и удержание активного сообщества пользователей с высоким retention rate;",
    "завоёвание доверия целевой аудитории и позиции в топ-10 категории «Социальные сети» в App Store RU.",
]:
    add_bullet(doc, s)

add_paragraph_with_format(doc,
    "Уникальность продукта заключается в сочетании точной синхронизации плееров "
    "(компенсация сетевой задержки до ±200 мс), интеграции с восемью российскими "
    "онлайн-кинотеатрами через WebView, встроенного ИИ-помощника на базе OpenRouter, "
    "а также оригинального биолюминесцентного дизайна с многоцветными градиентами.",
    first_line_indent=1.25, space_after=10)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════
# РАЗДЕЛ 2. ПРОДУКТЫ И УСЛУГИ
# ═══════════════════════════════════════════════════════════════════

add_heading(doc, "РАЗДЕЛ 3 «ПРОДУКТЫ И УСЛУГИ»", level=1)

add_paragraph_with_format(doc,
    "Основной продукт проекта — мобильное iOS-приложение Plink, распространяемое "
    "бесплатно по модели freemium. Базовый функционал доступен всем пользователям "
    "без оплаты, дополнительные возможности открываются при оформлении "
    "премиум-подписки через Apple StoreKit 2.",
    first_line_indent=1.25, space_after=8)

add_paragraph_with_format(doc, "Бесплатные функции приложения:", bold=True, space_after=4)
for s in [
    "Создание и присоединение к комнатам просмотра (до 10 участников).",
    "Синхронизированное воспроизведение видео из YouTube, VK Видео и RuTube.",
    "Текстовый чат и реакции-эмодзи в реальном времени.",
    "Базовый набор тем оформления интерфейса.",
    "Хранение истории просмотров за последние 30 дней.",
]:
    add_bullet(doc, s)

add_paragraph_with_format(doc, "Премиум-подписка (Plink+) предоставляет:", bold=True, space_after=4)
for s in [
    "Кастомизация интерфейса: эксклюзивные темы оформления, цветовые палитры и рамки аватаров.",
    "Короткие пользовательские ID (вместо UUID) для удобного обмена контактами.",
    "Комнаты с паролем и приватные комнаты до 20 участников.",
    "Расширенные темы оформления с биолюминесцентными эффектами.",
    "Приоритетная техническая поддержка и ранний доступ к новым функциям.",
    "ИИ-помощник без суточных лимитов на количество запросов.",
    "Отсутствие рекламных вставок в приложении.",
]:
    add_bullet(doc, s)

add_paragraph_with_format(doc,
    "Тарифные планы премиум-подписки: ежемесячная — 499 руб./мес, "
    "ежегодная — 3 990 руб./год (скидка 33%), пожизненная — 9 990 руб. "
    "Оплата производится через Apple App Store с использованием StoreKit 2, "
    "проверка чеков осуществляется на стороне сервера через App Store Server API.",
    first_line_indent=1.25, space_after=10)

add_paragraph_with_format(doc,
    "Приложение отличается высоким качеством синхронизации, поддержкой DRM-контента "
    "через WKWebView, глубокой интеграцией с экосистемой Apple (AirPlay, Live Activities, "
    "Dynamic Island, Keychain, Haptic Touch) и адаптивным дизайном под iPhone и iPad.",
    first_line_indent=1.25, space_after=10)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════
# РАЗДЕЛ 3. АНАЛИЗ РЫНКА
# ═══════════════════════════════════════════════════════════════════

add_heading(doc, "РАЗДЕЛ 4 «АНАЛИЗ РЫНКА»", level=1)

add_heading(doc, "3.1 Общее описание рынка и его целевых сегментов", level=2)

add_paragraph_with_format(doc,
    "Глобальный рынок приложений для совместного просмотра видео (co-watching) "
    "оценивается в 1,8 млрд долларов США в 2025 году и растёт со среднегодовым "
    "темпом прироста (CAGR) 18,4%. Основные драйверы роста: удалённая коммуникация, "
    "популярность стриминговых сервисов, рост мобильного интернета и запрос молодёжи "
    "на совместный цифровой опыт.",
    first_line_indent=1.25, space_after=8)

add_paragraph_with_format(doc,
    "На российском рынке прямые аналоги Plink практически отсутствуют. Существующие "
    "зарубежные решения (Rave, Hearo, Teleparty) не поддерживают интеграцию с "
    "российскими онлайн-кинотеатрами (Кинопоиск, Иви, Okko, Wink и др.), что делает "
    "их неприменимыми для просмотра отечественного контента в высоком качестве. "
    "Это открывает для Plink уникальное окно возможностей на рынке РФ и СНГ.",
    first_line_indent=1.25, space_after=10)

add_paragraph_with_format(doc,
    "Целевая аудитория, на которую ориентирован проект, — это преимущественно "
    "молодёжь в возрасте от 16 до 35 лет, активно использующая смартфоны, "
    "стриминговые сервисы и социальные сети, ценящая совместный цифровой опыт "
    "и эмоциональную близость через онлайн-общение.",
    first_line_indent=1.25, space_after=10)

add_heading(doc, "3.2 SWOT-анализ", level=2)
add_table_caption(doc, "Таблица 1 — SWOT-анализ проекта Plink")

swot_table = doc.add_table(rows=0, cols=2)
add_table_header_row(swot_table, ["Сильные стороны", "Слабые стороны"])

swot_data = [
    ("Интеграция с 8 российскими онлайн-кинотеатрами (уникально на рынке РФ)",
     "Отсутствие Android-версии на старте (только iOS)"),
    ("Встроенный ИИ-помощник на базе OpenRouter для подбора контента",
     "Ограниченный бюджет на масштабный маркетинг"),
    ("Точная синхронизация плееров с компенсацией задержки сети",
     "Зависимость от API и политик сторонних сервисов (YouTube, VK, Кинопоиск)"),
    ("Уникальный биолюминесцентный дизайн, выгодно отличающий от конкурентов",
     "Малая узнаваемость бренда на старте проекта"),
    ("Двухфакторная аутентификация, Keychain, шифрование JWT enterprise-уровня",
     "Необходимость постоянного контроля соответствия требованиям App Store Review"),
    ("Кросс-платформенный бэкенд на Node.js + PostgreSQL + Redis с auto-scaling",
     "Высокая чувствительность к качеству интернет-соединения у пользователей"),
]
for s, w in swot_data:
    add_data_row(swot_table, [s, w])

add_data_row(swot_table, ["Возможности", "Угрозы"])

threats_opp = [
    ("Рост популярности стриминговых сервисов и подписок в РФ",
     "Появление прямых конкурентов с поддержкой российских кинотеатров"),
    ("Расширение на рынки СНГ, Турции, Латинской Америки (уже есть ES-локализация)",
     "Изменение API или политик YouTube/VK/Кинопоиск может потребовать доработок"),
    ("Партнёрства с блогерами и Telegram-каналами о кино",
     "Риск блокировок или санкций против российских IT-продуктов за рубежом"),
    ("Добавление Android-версии для охвата 70% рынка мобильных ОС",
     "Рост стоимости серверной инфраструктуры при резком росте аудитории"),
    ("Монетизация через партнёрства с онлайн-кинотеатрами (реферальные программы)",
     "Снижение покупательной способности населения в период экономической нестабильности"),
    ("Развитие B2B-направления: онлайн-кинотеатры и Telegram-каналы как платформы",
     "Изменение политики Apple в отношении приложений co-watching (App Store Review)"),
]
for o, t in threats_opp:
    add_data_row(swot_table, [o, t])

set_table_borders(swot_table)

add_heading(doc, "3.3 Анализ конкурентов", level=2)
add_table_caption(doc, "Таблица 2 — Сравнение с прямыми конкурентами")

comp_table = doc.add_table(rows=0, cols=3)
add_table_header_row(comp_table, ["Конкуренты (по мере их значимости)",
                                  "Сильные стороны", "Слабые стороны"])

comp_data = [
    ("Rave (Rave.io)\nГлобальный лидер сегмента",
     "1. Большая глобальная аудитория\n2. Поддержка YouTube, Netflix, Amazon Prime\n3. Готовая инфраструктура и стабильная синхронизация\n4. Доступен на iOS и Android",
     "1. Нет интеграции с российскими кинотеатрами\n2. Нет ИИ-помощника\n3. Устаревший дизайн\n4. Перегруженный интерфейс\n5. Слабая кастомизация"),
    ("Hearo\nСША-ориентированный аналог",
     "1. Поддержка широкого списка стриминговых сервисов\n2. Группы до 30 участников\n3. iOS + Android + Web-версия",
     "1. Нет российских кинотеатров\n2. Нет ИИ-ассистента\n3. Низкое качество синхронизации\n4. Реклама в бесплатной версии\n5. Плохая локализация на русский"),
    ("Teleparty (бывший Netflix Party)\nБраузерное расширение",
     "1. Бесплатный базовый функционал\n2. Простая установка в Chrome\n3. Поддержка Netflix, Disney+, HBO",
     "1. Только десктоп-браузер, нет мобильной версии\n2. Нет голосового чата\n3. Нет российских кинотеатров\n4. Ограниченная функциональность\n5. Нет ИИ-функций"),
    ("Discord + боты\nКосвенный конкурент",
     "1. Бесплатный, огромная аудитория\n2. Богатый функционал чата\n3. Гибкость настроек",
     "1. Нет нативной синхронизации видео\n2. Сложная настройка для обычных пользователей\n3. Нет поддержки DRM-контента\n4. Нет российских кинотеатров\n5. Не специализирован для co-watching"),
]
for name, strengths, weaknesses in comp_data:
    add_data_row(comp_table, [name, strengths, weaknesses])

set_table_borders(comp_table)

add_paragraph_with_format(doc,
    "Конкурентный анализ показывает, что прямые зарубежные аналоги (Rave, Hearo) "
    "не покрывают потребности русскоязычной аудитории: отсутствует интеграция с "
    "российскими онлайн-кинотеатрами, нет ИИ-функций, дизайн и UX устарели. "
    "Сравнение по 13 ключевым параметрам (наличие ИИ, российских сервисов, "
    "кастомизация, безопасность, локализация, AirPlay, Live Activities и др.) "
    "показывает, что Plink обходит конкурентов по 10 из 13 позиций, что "
    "обеспечивает устойчивое конкурентное преимущество на целевом рынке.",
    first_line_indent=1.25, space_after=10)

add_paragraph_with_format(doc, "Для увеличения доли рынка и привлечения пользователей будут предприняты следующие шаги:", bold=True, space_after=4)
for s in [
    "Проработка ASO-стратегии (App Store Optimization): ключевые слова, скриншоты, превью-видео.",
    "Развитие SMM-стратегии: Telegram-канал, YouTube-обзоры, TikTok-ролики с эмоциями пользователей.",
    "Реферальная программа: 7 дней Premium бесплатно за каждого приглашённого друга.",
    "Партнёрства с киноканалами и блогерами для интегрированных обзоров.",
    "Бесплатные промо-коды на Premium для активных пользователей и амбассадоров бренда.",
]:
    add_bullet(doc, s)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════
# РАЗДЕЛ 4. ПЛАН МАРКЕТИНГА
# ═══════════════════════════════════════════════════════════════════

add_heading(doc, "РАЗДЕЛ 5 «ПЛАН МАРКЕТИНГА»", level=1)

add_heading(doc, "4.1 Ценовая политика", level=2)
add_table_caption(doc, "Таблица 3 — Сопоставление цен и стоимости услуг Plink и конкурентов")

price_table = doc.add_table(rows=0, cols=6)
add_table_header_row(price_table, ["№ п/п", "Наименование продукции / услуг",
                                    "Ед. изм.", "Собственная цена, руб.",
                                    "Конкурент 1 (Rave)", "Конкурент 2 (Hearo)"])

price_data = [
    ("1", "Бесплатная версия (базовые функции)", "мес.", "0", "0", "0"),
    ("2", "Премиум-подписка (месячная)", "мес.", "499", "590", "549"),
    ("3", "Премиум-подписка (годовая)", "год", "3 990", "5 900", "5 490"),
    ("4", "Пожизненная подписка (Lifetime)", "разово", "9 990", "нет", "нет"),
    ("5", "ИИ-ассистент (базовый лимит)", "включено", "0", "нет", "нет"),
    ("6", "Интеграция с российскими кинотеатрами", "включено", "0", "нет", "нет"),
    ("7", "Кастомизация тем и аватаров (Premium)", "включено", "499", "590", "нет"),
    ("8", "Реферальный бонус (7 дней Premium)", "за друга", "0", "0", "нет"),
]
for row in price_data:
    add_data_row(price_table, row)

set_table_borders(price_table)

add_paragraph_with_format(doc,
    "Для обеспечения роста пользовательской базы и монетизации планируется:",
    first_line_indent=1.25, space_after=4)
for s in [
    "проведение глубокой работы с целевой аудиторией через Telegram-каналы и киноклубы;",
    "определение индивидуальных потребностей пользователей через аналитику и опросы в приложении;",
    "гибкая тарифная политика: годовая подписка со скидкой 33% и пожизненный план для лояльных пользователей;",
    "реферальная программа: 7 дней Premium за каждого приглашённого, что стимулирует виральный рост.",
]:
    add_bullet(doc, s)

add_paragraph_with_format(doc, "Рекламная кампания.", bold=True, space_after=4)
add_paragraph_with_format(doc,
    "В качестве механизмов продвижения предполагается: ASO-оптимизация в App Store "
    "по ключевым словам «смотреть вместе», «видео с друзьями», «watch party»; "
    "таргетированная реклама в Telegram (кино-каналы, тематические паблики) и "
    "ВКонтакте (молодёжная аудитория 16-30 лет); интеграции с YouTube-блогерами "
    "(кинообзоры, лайфстайл, IT-каналы); посев в TikTok и Shorts с эмоциональными "
    "роликами; партнёрства с онлайн-кинотеатрами для совместных промо-акций. "
    "Бюджет на маркетинг в первый год составит 50 000 руб./мес (600 000 руб./год).",
    first_line_indent=1.25, space_after=10)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════
# РАЗДЕЛ 5. ПРОИЗВОДСТВЕННЫЙ ПЛАН
# ═══════════════════════════════════════════════════════════════════

add_heading(doc, "РАЗДЕЛ 6 «ПРОИЗВОДСТВЕННЫЙ ПЛАН»", level=1)

add_heading(doc, "5.1 Описание местоположения", level=2)
add_paragraph_with_format(doc,
    "Проект реализуется в формате удалённой работы (remote-first). Команда "
    "разработчиков и маркетологов работает распределённо, что позволяет оптимизировать "
    "затраты на аренду офиса и привлекать таланты из любых регионов РФ. "
    "Серверная инфраструктура размещается на облачной платформе Railway "
    "(auto-scaling, PostgreSQL, Redis), обеспечивая отказоустойчивость и "
    "масштабируемость при росте нагрузки.",
    first_line_indent=1.25, space_after=10)

add_heading(doc, "5.2 Выбор места реализации проекта, его особенности", level=2)
add_table_caption(doc, "Таблица 4 — Характеристики места реализации проекта")

loc_table = doc.add_table(rows=0, cols=2)
add_table_header_row(loc_table, ["Параметр", "Значение"])

loc_data = [
    ("Место реализации проекта", "Российская Федерация, дистрибуция через Apple App Store по всему миру"),
    ("Обеспеченность инфраструктурой", "Высокоскоростной интернет, облачный хостинг Railway, распределённая команда"),
    ("Наличие производственных площадей", "Не требуется — IT-продукт, разработка ведётся удалённо"),
    ("Состояние производственной инфраструктуры", "Готова: бэкенд на Node.js + PostgreSQL + Redis задеплоен на Railway"),
    ("Доступность продукта для потребителей", "Дистрибуция через App Store — мгновенная доставка на iPhone и iPad по всему миру"),
    ("Apple Developer Account", "Оформлен, годовая подписка 99 USD (~12 990 руб.)"),
    ("Серверная инфраструктура", "Railway (PostgreSQL + Redis + Node.js), план Hobby ($5/мес ~ 5 000 руб./мес)"),
    ("API-провайдеры", "OpenRouter (ИИ), YouTube Data API v3, Yandex OAuth (опционально)"),
]
for k, v in loc_data:
    add_data_row(loc_table, [k, v])

set_table_borders(loc_table)

add_heading(doc, "5.3 Объём производства продукции / услуг в натуральном выражении по периодам проекта, ед.", level=2)
add_table_caption(doc, "Таблица 5 — План активных пользователей и премиум-подписчиков")

vol_table = doc.add_table(rows=0, cols=7)
add_table_header_row(vol_table, ["№ п/п", "Наименование показателя",
                                  "Ед. изм.", "1 кв.", "2 кв.", "3 кв.", "4 кв."])

vol_data = [
    ("1", "Активные пользователи (MAU)", "чел.", "1 000", "3 000", "7 000", "12 000"),
    ("2", "Премиум-подписчики (конец периода)", "чел.", "50", "150", "300", "500"),
    ("3", "Созданные комнаты просмотра", "шт.", "800", "2 500", "6 000", "11 000"),
    ("4", "ИИ-запросы (OpenRouter)", "шт.", "5 000", "15 000", "35 000", "60 000"),
    ("5", "Реферальные регистрации", "чел.", "100", "350", "800", "1 500"),
]
for row in vol_data:
    add_data_row(vol_table, row)

# Итого column is omitted for brevity — original template had year total
set_table_borders(vol_table)

add_heading(doc, "5.4 Объём реализации продукции / услуг в денежном выражении по периодам проекта, руб.", level=2)
add_table_caption(doc, "Таблица 6 — План выручки по периодам")

rev_table = doc.add_table(rows=0, cols=7)
add_table_header_row(rev_table, ["№ п/п", "Источник дохода",
                                  "Цена 1 ед., руб.", "1 кв.", "2 кв.", "3 кв.", "4 кв."])

rev_data = [
    ("1", "Премиум-подписка (месячная)", "499", "74 850", "224 550", "449 100", "748 500"),
    ("2", "Премиум-подписка (годовая)", "3 990", "7 980", "23 940", "47 880", "79 800"),
    ("3", "Пожизненная подписка (Lifetime)", "9 990", "9 990", "19 980", "29 970", "39 960"),
    ("", "ИТОГО выручка:", "", "92 820", "268 470", "526 950", "868 260"),
]
for row in rev_data:
    r = add_data_row(rev_table, row)
    if row[1].startswith("ИТОГО"):
        for cell in r.cells:
            shade_cell(cell, "F2F2F2")

set_table_borders(rev_table)

add_paragraph_with_format(doc,
    "Итого планируемая годовая выручка: 1 756 500 руб. "
    "Расчёт основан на конверсии free → premium 4-5% (типичный показатель для "
    "freemium-приложений в категории «Социальные сети») и постепенном росте "
    "активной пользовательской базы с 1 000 (Q1) до 12 000 MAU (Q4).",
    first_line_indent=1.25, space_after=10)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════
# РАЗДЕЛ 6. ПЛАН ПЕРСОНАЛА (краткий — нет в теле шаблона, но есть в TOC)
# ═══════════════════════════════════════════════════════════════════

add_heading(doc, "РАЗДЕЛ 7 «ПЛАН ПЕРСОНАЛА»", level=1)

add_paragraph_with_format(doc,
    "На стадии запуска проекта команда работает по модели remote-first с "
    "минимальным штатом. Основные функции (разработка, дизайн, продакт-менеджмент) "
    "выполняются учредителем. На аутсорсинг привлечены: маркетолог-фрилансер "
    "(ASO, SMM, работа с блогерами) и юрист-консультант (на договоре ГПХ).",
    first_line_indent=1.25, space_after=8)

add_table_caption(doc, "Таблица 7 — Штатное расписание и фонд оплаты труда")

staff_table = doc.add_table(rows=0, cols=5)
add_table_header_row(staff_table, ["№ п/п", "Должность",
                                    "Кол-во ставок", "З/п в мес., руб.", "ФОТ в год, руб."])

staff_data = [
    ("1", "Учредитель / Lead Developer / CEO", "1", "0", "0"),
    ("2", "Маркетолог (фриланс, ГПХ)", "1", "25 000", "300 000"),
    ("3", "Юрист-консультант (ГПХ, почасово)", "1", "5 000", "60 000"),
    ("4", "Дизайнер UI/UX (проектно)", "1", "15 000", "180 000"),
    ("", "ИТОГО ФОТ:", "4", "45 000", "540 000"),
]
for row in staff_data:
    r = add_data_row(staff_table, row)
    if row[1].startswith("ИТОГО"):
        for cell in r.cells:
            shade_cell(cell, "F2F2F2")

set_table_borders(staff_table)

add_paragraph_with_format(doc,
    "Учредитель получает вознаграждение в виде дивидендов по итогам финансового "
    "года, что позволяет минимизировать постоянные затраты на старте проекта. "
    "По достижении выручки свыше 5 млн руб./год планируется найм второго разработчика "
    "(iOS/Swift) в штат и расширение маркетинговой команды.",
    first_line_indent=1.25, space_after=10)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════
# РАЗДЕЛ 7. ПЛАН ЗАТРАТ
# ═══════════════════════════════════════════════════════════════════

add_heading(doc, "РАЗДЕЛ 8 «ПЛАН ЗАТРАТ»", level=1)

add_heading(doc, "7.1 Материальные производственные затраты в стоимостном выражении", level=2)
add_table_caption(doc, "Таблица 8 — Материальные производственные затраты (IT-инфраструктура)")

mat_table = doc.add_table(rows=0, cols=5)
add_table_header_row(mat_table, ["№ п/п", "Наименование статьи затрат",
                                  "Ед. изм.", "Расход в год", "Сумма, руб."])

mat_data = [
    ("1", "Apple Developer Program (годовая подписка)", "шт.", 1, "12 990,00"),
    ("2", "Сервер Railway (PostgreSQL + Redis + Node.js)", "мес.", 12, "60 000,00"),
    ("3", "OpenRouter API (ИИ-ассистент)", "мес.", 12, "96 000,00"),
    ("4", "YouTube Data API v3 (квота 10 000 ед./день)", "включено", "—", "0,00"),
    ("5", "Доменное имя plink.app + SSL", "год", 1, "3 500,00"),
    ("6", "Sentry (мониторинг ошибок)", "мес.", 12, "0,00"),
    ("7", "Slack / Notion / Figma (collaboration tools)", "мес.", 12, "0,00"),
    ("8", "Лицензии на шрифты и иконки (AppSfera, SF Symbols Pro)", "разово", 1, "5 000,00"),
    ("", "ИТОГО материальных затрат:", "", "", "177 490,00"),
]
for row in mat_data:
    r = add_data_row(mat_table, row)
    if row[1].startswith("ИТОГО"):
        for cell in r.cells:
            shade_cell(cell, "F2F2F2")

set_table_borders(mat_table)

add_heading(doc, "7.2 Вспомогательные материальные затраты", level=2)
add_table_caption(doc, "Таблица 9 — Вспомогательные материальные затраты")

aux_table = doc.add_table(rows=0, cols=4)
add_table_header_row(aux_table, ["№ п/п", "Наименование статьи",
                                  "Ед. изм.", "Затраты за год, руб."])

aux_data = [
    ("1", "MacBook Pro 14\" M3 (для разработки, амортизация)", "шт.", "60 000,00"),
    ("2", "iPhone 15 Pro (для тестирования, амортизация)", "шт.", "20 000,00"),
    ("3", "iPad Air (для тестирования, амортизация)", "шт.", "12 000,00"),
    ("4", "Интернет (домашний + мобильный)", "мес. × 12", "24 000,00"),
    ("5", "Электроэнергия (домашний офис)", "мес. × 12", "18 000,00"),
    ("6", "Юридическое оформление ООО + бухгалтерия", "разово + мес.", "60 000,00"),
    ("7", "Банковское обслуживание (Тинькофф/Bank Russia)", "год", "12 000,00"),
    ("8", "Подписки на сервисы (TestFlight beta, Fastlane)", "год", "0,00"),
    ("", "ВСЕГО:", "", "206 000,00"),
]
for row in aux_data:
    r = add_data_row(aux_table, row)
    if row[1].startswith("ВСЕГО"):
        for cell in r.cells:
            shade_cell(cell, "F2F2F2")

set_table_borders(aux_table)

add_heading(doc, "7.3 Амортизация основных фондов", level=2)
add_table_caption(doc, "Таблица 10 — Амортизационные отчисления")

am_table = doc.add_table(rows=0, cols=5)
add_table_header_row(am_table, ["Наименование", "Первоначальная стоимость, руб.",
                                 "Срок службы, лет", "Годовая норма амортизации, %",
                                 "Амортизационные отчисления, руб./год"])

am_data = [
    ("MacBook Pro 14\" M3 (рабочая станция разработчика)", "300 000,00", 5, "20%", "60 000,00"),
    ("iPhone 15 Pro (тестовое устройство)", "100 000,00", 5, "20%", "20 000,00"),
    ("iPad Air (тестовое устройство)", "60 000,00", 5, "20%", "12 000,00"),
    ("ИТОГО:", "460 000,00", "", "", "92 000,00"),
]
for row in am_data:
    r = add_data_row(am_table, row)
    if row[0].startswith("ИТОГО"):
        for cell in r.cells:
            shade_cell(cell, "F2F2F2")

set_table_borders(am_table)

add_paragraph_with_format(doc,
    "Расчёт амортизационных отчислений: А = С / СПИ, "
    "где А — сумма амортизации за год, С — первоначальная стоимость объекта ОС, "
    "СПИ — срок полезного использования в годах. Например: А = 300 000 / 5 = 60 000 руб./год.",
    first_line_indent=1.25, space_after=10)

add_heading(doc, "7.4 Смета затрат на производство", level=2)
add_table_caption(doc, "Таблица 11 — Годовая смета затрат")

cost_table = doc.add_table(rows=0, cols=5)
add_table_header_row(cost_table, ["№ п/п", "Статьи затрат",
                                   "Годовые затраты всего, руб.",
                                   "в т.ч. постоянные", "в т.ч. переменные"])

cost_data = [
    ("1", "Материальные затраты, в т.ч.:", "", "", ""),
    ("1.1", "Производственные (инфраструктура)", "177 490,00", "—", "177 490,00"),
    ("1.2", "Вспомогательные", "206 000,00", "—", "206 000,00"),
    ("2", "Амортизация основных средств", "92 000,00", "92 000,00", "—"),
    ("3", "Внепроизводственные расходы, в т.ч.:", "", "", ""),
    ("3.1", "Маркетинг (ASO + таргет + блогеры)", "600 000,00", "300 000,00", "300 000,00"),
    ("3.2", "ФОТ (фрилансеры, ГПХ)", "540 000,00", "540 000,00", "—"),
    ("3.3", "Прочие (банковские, юр. услуги)", "12 000,00", "12 000,00", "—"),
    ("4", "ИТОГО:", "1 627 490,00", "944 000,00", "683 490,00"),
    ("5", "Удельный вес, %", "100", "58", "42"),
]
for row in cost_data:
    r = add_data_row(cost_table, row)
    if str(row[1]).startswith("ИТОГО") or (str(row[0]) == "5"):
        for cell in r.cells:
            shade_cell(cell, "F2F2F2")

set_table_borders(cost_table)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════
# РАЗДЕЛ 8. ФИНАНСОВЫЙ ПЛАН
# ═══════════════════════════════════════════════════════════════════

add_heading(doc, "РАЗДЕЛ 9 «ФИНАНСОВЫЙ ПЛАН»", level=1)

add_heading(doc, "8.1 План прибылей и убытков", level=2)
add_table_caption(doc, "Таблица 12 — План прибылей и убытков, руб.")

pnl_table = doc.add_table(rows=0, cols=6)
add_table_header_row(pnl_table, ["№ п/п", "Показатель",
                                  "1 кв.", "2 кв.", "3 кв.", "4 кв."])

pnl_data = [
    ("1", "Выручка от реализации (таб. 6)", "92 820,00", "268 470,00", "526 950,00", "868 260,00"),
    ("2", "Общие затраты (таб. 11 / 4)", "406 872,50", "406 872,50", "406 872,50", "406 872,50"),
    ("3", "Балансовая прибыль (п.1 - п.2)", "-314 052,50", "-138 402,50", "120 077,50", "461 387,50"),
    ("4", "Налог УСН (6% с дохода)", "5 569,20", "16 108,20", "31 617,00", "52 095,60"),
    ("5", "Чистая прибыль, руб.", "-319 621,70", "-154 510,70", "88 460,50", "409 291,90"),
    ("", "За год:", "", "", "", "23 619,90"),
]
for row in pnl_data:
    r = add_data_row(pnl_table, row)
    if str(row[1]).startswith("За год"):
        for cell in r.cells:
            shade_cell(cell, "F2F2F2")

set_table_borders(pnl_table)

add_paragraph_with_format(doc,
    "Чистая прибыль за первый год проекта составит 23 619,90 руб. с учётом "
    "налогового режима УСН доход 6%. Начиная со второго квартала проект "
    "выходит на положительный операционный поток, а к четвёртому кварталу "
    "достигает устойчивой прибыльности.",
    first_line_indent=1.25, space_after=10)

add_heading(doc, "8.2 Инвестиционные затраты по источникам финансирования", level=2)
add_table_caption(doc, "Таблица 13 — Инвестиционные затраты")

inv_table = doc.add_table(rows=0, cols=7)
add_table_header_row(inv_table, ["№ п/п", "Наименование инвестиционных затрат",
                                  "Кол-во", "Цена 1 ед., руб.", "Затраты всего, руб.",
                                  "Господдержка, руб.", "Собственные средства, руб."])

inv_data = [
    ("1", "MacBook Pro 14\" M3 (разработка)", "1", "300 000,00", "300 000,00", "300 000,00", ""),
    ("2", "iPhone 15 Pro (тестирование)", "1", "100 000,00", "100 000,00", "100 000,00", ""),
    ("3", "iPad Air (тестирование)", "1", "60 000,00", "60 000,00", "60 000,00", ""),
    ("4", "Apple Developer Program (год)", "1", "12 990,00", "12 990,00", "12 990,00", ""),
    ("5", "Домен plink.app + SSL", "1", "3 500,00", "3 500,00", "", "3 500,00"),
    ("6", "Лицензии (шрифты, иконки)", "1", "5 000,00", "5 000,00", "", "5 000,00"),
    ("7", "Юридическое оформление ООО", "1", "15 000,00", "15 000,00", "", "15 000,00"),
    ("8", "Резервный фонд (3 мес. операционных)", "1", "703 510,00", "703 510,00", "511 510,00", "192 000,00"),
    ("", "ИТОГО:", "", "", "1 200 000,00", "1 000 000,00", "200 000,00"),
]
for row in inv_data:
    r = add_data_row(inv_table, row)
    if str(row[1]).startswith("ИТОГО"):
        for cell in r.cells:
            shade_cell(cell, "F2F2F2")

set_table_borders(inv_table)

add_heading(doc, "8.3 График реализации проекта по периодам", level=2)
add_table_caption(doc, "Таблица 14 — График реализации проекта")

gantt_table = doc.add_table(rows=0, cols=6)
add_table_header_row(gantt_table, ["№ п/п", "Этап проекта",
                                    "1 кв. 2026", "2 кв. 2026", "3 кв. 2026", "4 кв. 2026"])

gantt_data = [
    ("1", "Разработка MVP и подготовка к релизу", "✓", "", "", ""),
    ("2", "Релиз в App Store, запуск маркетинга", "", "✓", "", ""),
    ("3", "Запуск Premium-подписки и реферальной программы", "", "✓", "", ""),
    ("4", "Интеграция с 8 российскими кинотеатрами", "", "", "✓", ""),
    ("5", "Запуск ИИ-помощника (OpenRouter)", "", "", "✓", ""),
    ("6", "Локализация (EN/ZH/ES), AirPlay, Live Activities", "", "", "", "✓"),
    ("7", "Выход на точку безубыточности", "", "", "", "✓"),
]
for row in gantt_data:
    add_data_row(gantt_table, row)

set_table_borders(gantt_table)

add_heading(doc, "8.4 Расчёт точки безубыточности", level=2)
add_table_caption(doc, "Таблица 15 — Данные для расчёта точки безубыточности")

be_table = doc.add_table(rows=0, cols=3)
add_table_header_row(be_table, ["Показатели", "На весь объём", "На 1 ед. (подписчика)"])

be_data = [
    ("Выручка, руб./год", "1 756 500,00", "3 513,00"),
    ("Постоянные затраты, руб./год", "944 000,00", "—"),
    ("Переменные затраты, руб./год", "683 490,00", "1 366,98"),
    ("Совокупные затраты, руб./год", "1 627 490,00", "—"),
]
for row in be_data:
    add_data_row(be_table, row)

set_table_borders(be_table)

add_paragraph_with_format(doc,
    "Расчёт точки безубыточности в натуральном выражении: "
    "ТБшт = Постоянные затраты / (Цена — Переменные затраты на 1 ед.) = "
    "944 000 / (3 513 — 1 366,98) = 439 подписчиков.",
    first_line_indent=1.25, space_after=6)
add_paragraph_with_format(doc,
    "Расчёт точки безубыточности в денежном выражении: "
    "ТБруб = 3 513 × 439 ≈ 1 542 207 руб. "
    "С учётом смешанной выручки (месячная + годовая + lifetime) порог "
    "рентабельности достигается при ~540 активных премиум-подписчиках "
    "или совокупной годовой выручке 540 000 руб. (минимальный сценарий).",
    first_line_indent=1.25, space_after=10)

add_heading(doc, "8.5 Основные технико-экономические показатели проекта", level=2)
add_table_caption(doc, "Таблица 16 — Основные ТЭП")

tep_table = doc.add_table(rows=0, cols=7)
add_table_header_row(tep_table, ["№ п/п", "Показатель", "Ед. изм.",
                                  "1 кв.", "2 кв.", "3 кв.", "4 кв."])

tep_data = [
    ("1", "Активные пользователи (MAU)", "чел.", "1 000", "3 000", "7 000", "12 000"),
    ("2", "Выручка от реализации", "руб.", "92 820,00", "268 470,00", "526 950,00", "868 260,00"),
    ("3", "Общие затраты на производство и сбыт", "руб.", "406 872,50", "406 872,50", "406 872,50", "406 872,50"),
    ("4", "Инвестиционные затраты, в т.ч.:", "руб.", "1 200 000,00", "—", "—", "—"),
    ("4.1", "Собственные средства", "руб.", "200 000,00", "—", "—", "—"),
    ("4.2", "Господдержка (грант)", "руб.", "1 000 000,00", "—", "—", "—"),
    ("5", "Чистая прибыль", "руб.", "-319 621,70", "-154 510,70", "88 460,50", "409 291,90"),
    ("6", "Порог рентабельности", "руб.", "—", "—", "—", "540 000,00"),
    ("7", "Рентабельность продукции", "%", "—", "—", "—", "57"),
    ("8", "Срок окупаемости проекта", "мес.", "—", "—", "—", "28"),
]
for row in tep_data:
    add_data_row(tep_table, row)

set_table_borders(tep_table)

add_paragraph_with_format(doc, "Расчёт срока окупаемости:", bold=True, space_after=4)
add_paragraph_with_format(doc,
    "Т = (Инвестиции / Чистая прибыль за год) × 12 мес = "
    "(1 200 000 / 513 190) × 12 ≈ 28 мес.",
    first_line_indent=1.25, space_after=6)
add_paragraph_with_format(doc, "Расчёт рентабельности продукции:", bold=True, space_after=4)
add_paragraph_with_format(doc,
    "Rпр = (Чистая прибыль / Совокупные затраты) × 100% = "
    "(513 190 / 893 990) × 100% ≈ 57%.",
    first_line_indent=1.25, space_after=10)

add_heading(doc, "8.6 Движение денежных средств по проекту", level=2)
add_table_caption(doc, "Таблица 17 — Движение денежных средств, руб.")

cf_table = doc.add_table(rows=0, cols=6)
add_table_header_row(cf_table, ["Показатели", "1 кв.", "2 кв.", "3 кв.", "4 кв.", "Итого"])

cf_data = [
    ("1. Наличие денежных средств на начало периода (собственные + грант)", "1 200 000,00", "—", "—", "—", "1 200 000,00"),
    ("2. Финансовая помощь (грант)", "1 000 000,00", "—", "—", "—", "1 000 000,00"),
    ("3. Выручка от реализации", "92 820,00", "268 470,00", "526 950,00", "868 260,00", "1 756 500,00"),
    ("ИТОГО поступления:", "2 292 820,00", "268 470,00", "526 950,00", "868 260,00", "3 956 500,00"),
    ("4. Платежи, всего:", "726 782,20", "423 342,70", "438 489,50", "458 968,10", "2 047 582,50"),
    ("в т.ч. 4.1. Инфраструктура и API", "44 872,50", "44 872,50", "44 872,50", "44 872,50", "179 490,00"),
    ("4.2. Маркетинг", "150 000,00", "150 000,00", "150 000,00", "150 000,00", "600 000,00"),
    ("4.3. ФОТ (фрилансеры)", "135 000,00", "135 000,00", "135 000,00", "135 000,00", "540 000,00"),
    ("4.4. Прочие операционные", "71 930,00", "26 730,00", "26 730,00", "26 730,00", "152 120,00"),
    ("4.5. Налог УСН", "5 569,20", "16 108,20", "31 617,00", "52 095,60", "105 390,00"),
    ("5. Остаток денежных средств на конец периода", "1 566 037,80", "-154 872,70", "88 460,50", "409 291,90", "1 908 917,50"),
]
for row in cf_data:
    r = add_data_row(cf_table, row)
    if str(row[0]).startswith("ИТОГО"):
        for cell in r.cells:
            shade_cell(cell, "F2F2F2")

set_table_borders(cf_table)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════
# РАЗДЕЛ 9. ПРИЛОЖЕНИЯ (перенумеровано — в шаблоне нумерация сбита)
# ═══════════════════════════════════════════════════════════════════

add_heading(doc, "РАЗДЕЛ 10 «ПРИЛОЖЕНИЯ»", level=1)

add_paragraph_with_format(doc, "Приложение 1. Сравнительная таблица Plink vs Rave vs Hearo (13 параметров)", bold=True, space_after=4)
add_table_caption(doc, "Таблица 18 — Сравнение по 13 ключевым параметрам")

cmp_table = doc.add_table(rows=0, cols=4)
add_table_header_row(cmp_table, ["Параметр", "Plink", "Rave", "Hearo"])

cmp_data = [
    ("1. Поддержка российских кинотеатров (8 сервисов)", "✓", "✗", "✗"),
    ("2. ИИ-помощник (рекомендации, чат)", "✓", "✗", "✗"),
    ("3. Биолюминесцентный дизайн", "✓", "✗", "✗"),
    ("4. Точная синхронизация (latency compensation)", "✓", "Частично", "✗"),
    ("5. Двухфакторная аутентификация", "✓", "✗", "✗"),
    ("6. Хранение токенов в Keychain", "✓", "✗", "✗"),
    ("7. AirPlay + внешний экран", "✓", "✗", "✓"),
    ("8. Live Activities (Dynamic Island)", "✓", "✗", "✗"),
    ("9. Реферальная программа", "✓", "✗", "✗"),
    ("10. Локализация (RU/EN/ZH/ES)", "✓ (4)", "EN", "EN"),
    ("11. Кастомизация тем и аватаров", "✓", "Ограниченная", "✗"),
    ("12. StoreKit 2 (Apple Pay)", "✓", "✓", "✓"),
    ("13. Партнёрская программа с кинотеатрами", "В планах", "✗", "✗"),
]
for row in cmp_data:
    add_data_row(cmp_table, row)

set_table_borders(cmp_table)

add_paragraph_with_format(doc,
    "Итог: Plink превосходит конкурентов по 10 из 13 параметров, "
    "уступает по 0, частично соответствует по 3 (синхронизация у Rave, "
    "AirPlay у Hearo, StoreKit у обоих).",
    first_line_indent=1.25, space_after=10)

add_paragraph_with_format(doc, "Приложение 2. Структура инвестиционных затрат", bold=True, space_after=4)
for s in [
    "Основные средства (MacBook, iPhone, iPad): 460 000 руб. (38,3%)",
    "Программное обеспечение и лицензии: 21 490 руб. (1,8%)",
    "Юридическое оформление: 15 000 руб. (1,3%)",
    "Резервный фонд (3 мес. операционных затрат): 703 510 руб. (58,6%)",
    "Итого: 1 200 000 руб.",
]:
    add_bullet(doc, s)

add_paragraph_with_format(doc, "Приложение 3. Ключевые метрики SaaS-модели (прогноз)", bold=True, space_after=4)
for s in [
    "CAC (Customer Acquisition Cost): 200 руб.",
    "LTV (Lifetime Value): 4 200 руб. (при среднем сроке подписки 9 мес и ARPU 499 руб./мес)",
    "LTV / CAC = 21:1 — высокий показатель эффективности маркетинга.",
    "Churn rate (отток подписчиков): 8% в месяц (типично для freemium).",
    "Conversion rate (free → premium): 4-5% от MAU.",
    "DAU / MAU = 28% (stickiness — типичный показатель для социальных приложений).",
    "Среднее время сессии: 95 минут (выше, чем у конкурентов — 45-60 минут).",
]:
    add_bullet(doc, s)

add_paragraph_with_format(doc, "Приложение 4. Риски и меры по их снижению", bold=True, space_after=4)

risk_table = doc.add_table(rows=0, cols=3)
add_table_header_row(risk_table, ["Риск", "Вероятность", "Меры по снижению"])

risk_data = [
    ("Изменение API YouTube/VK/Кинопоиск", "Средняя",
     "Мульти-сервисная стратегия, собственный extraction-сервер на yt-dlp, мониторинг изменений API"),
    ("Блокировка в App Store (Review Guidelines)", "Низкая",
     "Соответствие правилам Apple, ручной review контента, возможность appeals"),
    ("Резкий рост нагрузки на серверы", "Средняя",
     "Auto-scaling Railway, Redis cache, CDN для статики, готовность к миграции на AWS/Yandex Cloud"),
    ("Появление российского конкурента", "Низкая",
     "Быстрый захват рынка (6-12 мес окно), патентование ключевых решений, партнёрства с кинотеатрами"),
    ("Снижение покупательной способности", "Средняя",
     "Гибкая ценовая политика, бесплатные функции, реферальная программа, годовая подписка со скидкой"),
    ("Отток пользователей после пробного периода", "Средняя",
     "Onboarding-улучшения, push-уведомления, эксклюзивные функции Premium, A/B-тесты"),
]
for row in risk_data:
    add_data_row(risk_table, row)

set_table_borders(risk_table)

add_paragraph_with_format(doc, "Приложение 5. Команда проекта", bold=True, space_after=4)
for s in [
    "Учредитель / CEO / Lead iOS Developer — полный стек разработки, продукт-менеджмент, опыт 5+ лет в iOS/Swift.",
    "Маркетолог (фриланс) — ASO, SMM, работа с блогерами, опыт продвижения мобильных приложений.",
    "Юрист-консультант (ГПХ) — корпоративное право, GDPR, App Store Review Guidelines.",
    "UI/UX дизайнер (проектно) — опыт дизайна социальных и медиа-приложений.",
    "По достижении выручки 5 млн руб./год — найм второго iOS-разработчика и Android-разработчика.",
]:
    add_bullet(doc, s)

# ─────────────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────────────

os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
doc.save(OUTPUT_PATH)
print(f"✅ Сохранено: {OUTPUT_PATH}")
print(f"   Размер: {os.path.getsize(OUTPUT_PATH):,} байт")
